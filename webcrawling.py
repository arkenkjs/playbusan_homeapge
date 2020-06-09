from urllib.request import urlopen
from bs4 import BeautifulSoup
import json
import os
from docx import Document


## python파일의 위치
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# DOC 오브젝트 생성
report = Document()
report.add_heading('2020년 01월 08일 인터넷에 기재된 교재 목록')

for page in range(1, 152):
    #html 내의 웹페이지 정보를 잘라서 bsObject에 가져온다.
    print('A'+str(page)+'B')
    url = 'https://www.aladin.co.kr/shop/wbrowse.aspx?BrowseTarget=List&ViewRowsCount=50&ViewType=Detail&PublishMonth=0&SortOrder=5&page='+\
            str(page)+'&Stockstatus=1&PublishDay=84&CID=131522&CustReviewRankStart=&CustReviewRankEnd=&CustReviewCountStart=&CustReviewCountEnd=&PriceFilterMin=&PriceFilterMax='
    html = urlopen(url).read()
    bsObject = BeautifulSoup(html, "html.parser") 

    # print(bsObject)#웹문서 전체가 출력됨
    # print(bsObject.head.title)#태그로 구성된 트리에서 타이틀만 출력

    #웹문서중에서 메타 데이터만 찾아서 content 속성을 가져옴
    # for meta in bsObject.head.find_all('meta'):
    #     print(meta.get('content'))

    #<meta content="The official home of the Python Programming Language" name="description"/>
    #위 meta 태그에서 name항목에 dscription의 content 부분만 출력
    #print(bsObject.head.find("meta",{"name":"description"}).get('content'))

    #웹내 모든 a태그에서 텍스트와 주소를 가져온다.
    # for link in bsObject.find_all('a'):
    #     print(link.text.strip(), link.get('href'))

    # 책의 상세 웹페이지 주소를 추출하여 리스트에 저장합니다.
    # book_page_urls = []
    # for cover in bsObject.find_all('a', {'class':'ss_f_g2'}):
    #     link = cover.select('a')[0].get('href')
    #     book_page_urls.append(link)

    #Myform > div:nth-child(2) > div:nth-child(1) > table > tbody > tr > td:nth-child(3) > table > tbody > tr:nth-child(1) > td:nth-child(1) > div:nth-child(1)
    book_list = bsObject.select('tr > td > div')
    for i, book in enumerate(book_list):
        # with open(os.path.join(BASE_DIR, 'result.txt'), 'w+', encoding='utf-8') as file:        
        #     file.write(book.text) 
        if i>49 and i<348:        
            report.add_paragraph(book.text)
    # print(bsObject.find_all('td > div', {'class':'detail'}))

report.save('result.docx')
    


